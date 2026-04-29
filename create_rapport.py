from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(0x1A, 0x37, 0x6C))
    # Ligne de soulignement via bordure bas
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1A376C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(0x27, 0x6F, 0xBF))
    return p

def body(text, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2 + level * 0.2)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True, color=(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A376C')
        tcPr.append(shd)

    # Données
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.paragraphs[0].clear()
            bold  = j == 0
            color_val = (0x2E, 0x86, 0x5A) if j == len(row) - 1 and "−" in str(val) else None
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=11, bold=bold, color=color_val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            # Alternance de couleur
            if i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F4F8')
                tcPr.append(shd)

    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


# ═══════════════════════════════════════════════════════════════════════════════
# TITRE
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prévision du niveau des nappes phréatiques\npar réseaux de neurones")
set_font(run, size=18, bold=True, color=(0x1A, 0x37, 0x6C))

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Marius — Master 1 — 2025/2026")
set_font(run, size=11, color=(0x55, 0x55, 0x55))


# ═══════════════════════════════════════════════════════════════════════════════
# 1. CONTEXTE ET DONNÉES
# ═══════════════════════════════════════════════════════════════════════════════

heading1("1. Contexte et données")

body("L'objectif est de prédire le niveau des nappes phréatiques (GWL, en mètres) "
     "à partir de variables environnementales mesurées mensuellement sur des puits australiens.")

heading2("Jeu de données")
bullet("2 510 fichiers CSV — un par puit — couvrant l'ensemble du territoire australien")
bullet("8 régions : VIC (1113 puits), NSW (687), WA (262), QLD (203), SA (198), TAS (19), NT (19), ACT (9)")
bullet("6 variables par puit : GWL (cible), Date, P (précipitations), T (température), "
       "ET (évapotranspiration), NDVI (indice de végétation)")
bullet("Fichier OUVRAGES.csv : identifiant, latitude, longitude, région de chaque puit")
body("Note : de nombreux puits présentent des valeurs manquantes (NaN) dans la colonne GWL, "
     "réduisant le nombre de puits exploitables selon les parties.")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. PARTIE 1 — VISUALISATION
# ═══════════════════════════════════════════════════════════════════════════════

heading1("2. Partie 1 — Visualisation")

heading2("Approche")
body("Une fonction générique plot_well(well_id) a été développée pour visualiser n'importe quel puit. "
     "Elle génère deux figures de 5 sous-graphiques chacune (un par variable) :")
bullet("Figure A — Valeurs brutes : courbes avec les unités réelles (m, mm/mois, °C, sans unité)")
bullet("Figure B — Valeurs normalisées : z-score appliqué colonne par colonne, "
       "permettant de comparer des variables d'unités différentes sur un même axe")

body("Formule de normalisation (z-score) :")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
run = p.add_run("x̃ = (x − μ) / σ       avec μ = moyenne, σ = écart-type de la séquence")
set_font(run, size=11)

body("La fonction plot_sample() sélectionne automatiquement n puits par région "
     "(via OUVRAGES.csv) et sauvegarde les graphiques au format PNG.")


# ═══════════════════════════════════════════════════════════════════════════════
# 3. PARTIE 2a — MODÈLE LSTM GÉNÉRIQUE
# ═══════════════════════════════════════════════════════════════════════════════

heading1("3. Partie 2a — Modèle LSTM générique")

heading2("Approche")
body("Un unique réseau LSTM est entraîné sur l'ensemble des puits australiens (modèle générique). "
     "Il apprend à prédire GWL à partir des variables environnementales passées.")

bullet("Découpage train/test : toute la chronique moins les 12 derniers mois pour l'entraînement, "
       "les 12 derniers mois pour le test (conforme aux consignes)")
bullet("Features d'entrée : P, T, ET, NDVI + latitude normalisée + longitude normalisée (6 variables)")
bullet("Cible : GWL au mois suivant")
bullet("Séquences glissantes de 12 mois (SEQ_LEN = 12)")
bullet("Normalisation z-score par puit et par variable, calculée sur les données d'entraînement uniquement")

heading2("Architecture")
bullet("Type : LSTM (Long Short-Term Memory) — réseau récurrent avec mémoire à long terme")
bullet("Configuration finale (issue du grid search) : 3 couches, hidden=128, dropout=0.2")
bullet("Couche de sortie : fully connected (128 → 1) — prédit une valeur scalaire GWL")
bullet("Optimiseur : Adam | Fonction de perte : MSE")

heading2("Grid Search — 108 combinaisons testées")
body("Recherche exhaustive sur la grille suivante, résultats sauvegardés dans grid_search_results.csv :")

add_table(
    ["Hyperparamètre", "Valeurs testées", "Meilleure valeur"],
    [
        ["hidden (neurones/couche)", "64, 128, 256", "128"],
        ["n_layers (couches LSTM)", "1, 2, 3", "3"],
        ["lr (learning rate)", "1e-3, 5e-4, 1e-4", "5e-4"],
        ["batch_size", "256, 512", "512"],
        ["epochs", "30, 50", "30"],
    ],
    col_widths=[5.5, 5.5, 4.5]
)

heading2("Optimisation GPU")
body("Le dataset complet a été préchargé en VRAM (RTX 5070 Ti — 16 Go) sous forme de tenseurs PyTorch. "
     "Le batching est effectué directement sur la GPU via torch.randperm(), "
     "éliminant tout transfert CPU→GPU pendant l'entraînement.")

heading2("Résultats — 854 puits")
add_table(
    ["Métrique", "Moyenne", "Écart-type", "Médiane"],
    [
        ["RMSE (m)", "0.805", "±0.968", "0.440"],
        ["MAE (m)",  "0.686", "±0.869", "0.356"],
    ],
    col_widths=[5, 4, 4, 4]
)
body("La médiane est privilégiée : l'écart-type élevé indique la présence de quelques puits "
     "très difficiles à prédire qui tirent la moyenne vers le haut.", space_after=2)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. PARTIE 2b — CHRONOS 2 (FOUNDATION MODEL)
# ═══════════════════════════════════════════════════════════════════════════════

heading1("4. Partie 2b — Chronos 2 (Foundation Model)")

heading2("Approche")
body("Utilisation du modèle pré-entraîné amazon/chronos-bolt-small issu de la librairie "
     "chronos-forecasting (v2.2.2). Aucun entraînement sur nos données.")

bullet("Modèle univarié : seul l'historique GWL est utilisé comme contexte (pas P, T, ET, NDVI)")
bullet("Prédiction zero-shot : le modèle prédit directement sans fine-tuning")
bullet("Contexte : toutes les valeurs GWL non-NaN avant les 12 derniers mois")
bullet("Prédiction probabiliste : N scénarios générés → médiane retenue comme prédiction finale")
bullet("Traitement par batch de 32 puits en parallèle sur GPU")
bullet("Critère d'inclusion : minimum 24 mois de GWL non-NaN dans la période d'entraînement")

heading2("Résultats — 1 174 puits")
add_table(
    ["Métrique", "Moyenne", "Écart-type", "Médiane"],
    [
        ["RMSE (m)", "0.673", "±0.973", "0.301"],
        ["MAE (m)",  "0.578", "±0.847", "0.248"],
    ],
    col_widths=[5, 4, 4, 4]
)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. COMPARAISON ET CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

heading1("5. Comparaison et conclusion")

add_table(
    ["Métrique", "LSTM (2a)", "Chronos 2 (2b)", "Gain"],
    [
        ["RMSE moyen (m)",  "0.805", "0.673", "−16 %"],
        ["MAE moyen (m)",   "0.686", "0.578", "−16 %"],
        ["RMSE médian (m)", "0.440", "0.301", "−32 %"],
        ["MAE médian (m)",  "0.356", "0.248", "−30 %"],
        ["Puits évalués",   "854",   "1 174", "—"],
    ],
    col_widths=[5.5, 3.5, 4, 3.5]
)

doc.add_paragraph()
body("Chronos 2 surpasse le LSTM sur l'ensemble des métriques, avec une réduction d'erreur "
     "de 30 à 32 % sur les médianes. Ce résultat est d'autant plus remarquable que :")
bullet("Chronos n'utilise que l'historique GWL (univarié), sans accès aux variables météo ni GPS")
bullet("Aucun entraînement sur les données australiennes n'a été effectué (zero-shot)")
bullet("Le LSTM dispose pourtant de 6 features d'entrée et de 30 epochs d'entraînement")

body("Explication : les nappes phréatiques présentent des patterns temporels réguliers "
     "(cycles saisonniers, tendances lentes) que Chronos 2 reconnaît grâce à son "
     "pré-entraînement sur des milliers de séries temporelles diverses. Le LSTM doit "
     "apprendre ces patterns depuis zéro avec un volume de données relatif limité.")

body("Limite de la comparaison : les ensembles de puits évalués diffèrent "
     "(854 pour le LSTM, 1174 pour Chronos) en raison de critères d'inclusion différents. "
     "Une comparaison stricte sur le même sous-ensemble de puits renforcerait la validité des conclusions.")


# ─── Sauvegarde ──────────────────────────────────────────────────────────────
OUTPUT = "rapport_nappes.docx"
doc.save(OUTPUT)
print(f"Rapport créé : {OUTPUT}")
